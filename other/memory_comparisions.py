import requests
from datetime import datetime, timedelta
import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import RGBColor

nodes_file_path = 'nodes.json'
with open(nodes_file_path, 'r') as file:
    nodes_data = json.load(file)
PROMETHEUS = "http://192.168.128.50:9090"
API_PATH = '/api/v1/query_range'


prev_ist_start_time = '2023-08-03 00:32'
prev_ist_end_time = '2023-08-03 10:32'


curr_ist_start_time = '2023-08-09 05:06' 
curr_ist_end_time = '2023-08-09 15:06'


queries = {"Host" : 'avg((uptycs_memory_used/uptycs_total_memory) * 100) by (host_name)',
           'Rule Engine' : "avg(uptycs_app_memory{app_name=~'.*ruleEngine.*'}) by (host_name)",
           'Osquery Ingestion' : "sum(uptycs_app_memory{app_name=~'osqueryIngestion'}) by (host_name)",
           "Kafka" : "avg(uptycs_app_memory{app_name=~'kafka'}) by (host_name)",
           "Trino" : "avg(uptycs_app_memory{app_name='trino'}) by (host_name)",
           "Tls" : "avg(uptycs_app_memory{app_name='tls'}) by (host_name)",
           "EventsDbIngestion" : "avg(uptycs_app_memory{app_name=~'eventsDbIngestion'}) by (host_name)",
           "Logger" : "sum(uptycs_app_memory{app_name=~'.*osqLogger-1.*'}) by (host_name)"
           }




def process(queries,ist_start_time,ist_end_time):
    final=dict()
    ist_time_format = '%Y-%m-%d %H:%M'
    start_time_utc = (datetime.strptime(ist_start_time, ist_time_format))
    end_time_utc = (datetime.strptime(ist_end_time, ist_time_format))
    stu = int(start_time_utc.timestamp())
    etu = int(end_time_utc.timestamp())
    for query in queries:
        final[query] = {}
        PARAMS = {
            'query': queries[query],
            'start': stu,
            'end': etu,
            'step':15
        }
        response = requests.get(PROMETHEUS + API_PATH, params=PARAMS)
        result = response.json()['data']['result']

        for res in result:
            hostname = res['metric']['host_name']
            values = [float(i[1]) for i in res['values']]   
            avg = sum(values) / len(values)
            final[query][hostname] = {"percentage":avg}
            if hostname+'v' in nodes_data:
                final[query][hostname]["GB"] = avg * float(nodes_data[hostname+'v']['ram']) / 100
            else:
                final[query][hostname]["GB"] = None
    return final


prev_final =  process(queries,prev_ist_start_time,prev_ist_end_time)
curr_final = process(queries,curr_ist_start_time,curr_ist_end_time)



# Rest of your code...

# Create a new Document
doc = Document()
doc.add_heading('Memory Usage Report', level=1)

def create_table(old_data , new_data):
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = True

    # Add table header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'prev build'
    header_cells[2].text = 'curr build'
    header_cells[3].text = 'Absolute diff(GB,%)'
    header_cells[4].text = 'Relative diff(%)'

    for query in new_data:
        for host_name in new_data[query]:
            row = table.add_row().cells
            row[0].text = f"Memory consumed by {query} {host_name}"

            try : 
                # row[1].text = f"{old_data[query][host_name]['GB']:.2f} gb ({old_data[query][host_name]['percentage']:.2f}%)"
                row[1].text = f"({old_data[query][host_name]['percentage']:.2f}%)"
            except Exception as e: 
                row[1].text = '-'

            # row[2].text = f"{new_data[query][host_name]['GB']:.2f} gb ({new_data[query][host_name]['percentage']:.2f}%)"
            row[2].text = f"({new_data[query][host_name]['percentage']:.2f}%)"


            try : 
                # diff = float(old_data[query][host_name]['GB']) - float(new_data[query][host_name]['GB'])
                percent_diff = float(old_data[query][host_name]['percentage']) - float(new_data[query][host_name]['percentage'])
                # difference_text = f"{abs(diff):.2f} gb ({abs(percent_diff):.2f}%)"
                difference_text = f"({abs(percent_diff):.2f}%)"
                row4_text = f"{abs(percent_diff) *100  / (float(old_data[query][host_name]['percentage'])) :.2f} %"
                if percent_diff<0:
                    run = row[3].add_paragraph().add_run(difference_text + "⬆️")
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red color

                    run = row[4].add_paragraph().add_run(row4_text+ "⬆️")
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                else:
                    run = row[3].add_paragraph().add_run(difference_text + "⬇️")
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                    run = row[4].add_paragraph().add_run(row4_text + "⬇️")
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
            except Exception as e:
                print(e)
                row[3].text = '-'
                row[4].text = '-'

                
                

                
               
                

           
create_table(prev_final , curr_final)
doc.save('memory_usage_report.docx')
