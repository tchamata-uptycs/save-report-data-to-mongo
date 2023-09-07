#takes a document, add a table in it(with data as the provided dict) and returns the same doc
from docx.shared import RGBColor
import os
import openpyxl
import json
from openpyxl.utils import get_column_letter
from openpyxl.styles import PatternFill, Font
from collections import defaultdict
import shutil,socket,paramiko
import concurrent.futures
import pymongo

def add_screenshots_to_docx(doc,directory_path, grafana_ids, table_ids):
    doc.add_heading("Charts", level=2)
    all_shots = os.listdir(directory_path)
    mapping = defaultdict(lambda: defaultdict(lambda: []))
    for shot in all_shots:
        panel_id, n, topic = shot.split('_')
        panel_id,n = int(panel_id),int(n)
        title, _ = topic.split('.')
        mapping[panel_id]['title'] = [title]
        mapping[panel_id]['images'].append(shot)
    
    for index, panel in enumerate(grafana_ids):
        if panel in table_ids:
            mapping[panel]['images'].sort(key=lambda s: int(s.split('_')[1]))
        if panel not in mapping:continue
        title = mapping[panel]['title'][0].capitalize()
        print(f'{index+1}:({panel}){title}')            
        doc.add_heading(title, level=3)
        for filename in mapping[panel]['images']:
            file_path = os.path.join(directory_path, filename)
            try:
                # img = PILImage.open(file_path)
                # img_width, img_height = img.size
                # aspect_ratio = img_height / img_width
                # width = 450
                # height = int(width * aspect_ratio)
                doc.add_picture(file_path)
            except Exception as e:
                print(f"Error processing image {filename}: {e}")
    try:
        shutil.rmtree(directory_path)
        print(f"Folder '{directory_path}' and its contents have been deleted successfully.")
    except Exception as e:
        print(f"Error deleting folder: {str(e)}")
    return doc

def add_table(doc,data_dict):
    title = data_dict["title"]
    header = data_dict["header"]
    body = data_dict["body"]
    doc.add_heading(title, level=3)
    cols = len(header)
    rows = len(body)
    table =doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    table.autofit = True
    header_cells = table.rows[0].cells
    for col in range(cols):
        header_cells[col].text = header[col]

    for row_num in range(rows):
        row = table.add_row().cells
        for col_num in range(len(body[row_num])):
            try:
                text,color = body[row_num][col_num]
            except:
                text = body[row_num][col_num]
                color="none"

            run = row[col_num].paragraphs[0].add_run(str(text))

            if color == "green":
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
            elif color=="red":
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
    return doc

def update_col_width(sheet,value_width,col):
    col_letter = get_column_letter(col)
    if col_letter not in sheet.column_dimensions:
        sheet.column_dimensions[col_letter].width = value_width
    else:
        current_width = sheet.column_dimensions[col_letter].width
        if value_width > current_width:
            sheet.column_dimensions[col_letter].width = value_width

def add_row(sheet,value,row_num,temp_col,found):
    for val in value:
        color=None
        try:
            num,color = val
        except:
            num=val
        try:
            num=float(num)
        except:pass
        
        new_cell = sheet.cell(row=row_num, column=temp_col, value=num)
        if not found:
            new_cell.fill = PatternFill(start_color="b8f1ff", end_color="b8f1ff", fill_type="solid")  # blue fill
        if color and color=="red":
            new_cell.fill = PatternFill(start_color="F0C0C0", end_color="F0C0C0", fill_type="solid")  # Red fill
        elif color and color=="green":
            new_cell.fill = PatternFill(start_color="C0F0CF", end_color="C0F0CF", fill_type="solid")  # Green fill
        temp_col+=1

        update_col_width(sheet,len(str(num)) , temp_col)
        update_col_width(sheet,len(str(num)) , temp_col-1)

    return sheet

def add_columns(sheet_name,data_dict,workbook,build):
    sheet = workbook[sheet_name]
    col = sheet.max_column +1 
    while(col>2 and sheet.cell(row=1,column=col-1).value is None):
        col-=1
        sheet.delete_cols(col)

    build_cell = sheet.cell(row=1, column=col, value=int(build))
    build_cell.font = Font(bold=True)
    metric_cell = sheet.cell(row=1, column=1, value="Metric")
    metric_cell.font = Font(bold=True)
    sheet.freeze_panes = sheet.cell(row=2, column=2)

    for lst in data_dict:
        key = lst[0]
        value = lst[1:]
        found=False
        row = sheet.max_row +1

        for row_num in range(2, row):
            cell_value = sheet.cell(row=row_num, column=1).value
            if cell_value is not None and cell_value.lower() == key.lower():
                found=True
                sheet = add_row(sheet,value,row_num,col,found)
                break  # Exit the loop after finding a match
        if not found:
            sheet.cell(row=row, column=1, value=key.lower())
            sheet = add_row(sheet,value,row,col,found)

    update_col_width(sheet,len(build)+1 , col)
    update_col_width(sheet,len(build)+1 , col-1)
    update_col_width(sheet,58 , 1)

def excel_update(sheets , prev_path, curr_path,build):
    if os.path.exists(curr_path):
        workbook = openpyxl.load_workbook(curr_path)
        missing_sheets = [sheet_name for sheet_name in sheets if sheet_name not in workbook.sheetnames]
        for sheet_name in missing_sheets:
            print(f"Added sheet {sheet_name}")
            workbook.create_sheet(sheet_name)
        print("Loaded existing excel and added new sheets")
    else:
        if not os.path.exists(prev_path):
            workbook = openpyxl.Workbook()
            default_sheet = workbook.active
            workbook.remove(default_sheet)
            for sheet in sheets:
                print(f"Added sheet {sheet}")
                sheet1 = workbook.create_sheet(sheet)
            print("created new excel sheet and added new sheets")
        else:
            workbook = openpyxl.load_workbook(prev_path)
            missing_sheets = [sheet_name for sheet_name in sheets if sheet_name not in workbook.sheetnames]
            for sheet_name in missing_sheets:
                print(f"Added sheet {sheet_name}")
                workbook.create_sheet(sheet_name)
            print("Loaded existing excel and added new sheets")
    for sheet in sheets:
        print(f"Processing sheet {sheet} ...")
        add_columns(sheet,sheets[sheet],workbook,build)
    
    workbook.save(curr_path)
    print("Updated Excel sheet succcessfully")


def add_load_details(doc,details):
    doc.add_heading("Load Details", level=2)
    for key, value in details.items():
        doc.add_paragraph(f"{' '.join(str(key).split('_')).title()}: {value}")
    return doc

def extract_node_detail(data,node_type,prom_con_obj):
    port=prom_con_obj.ssh_port
    username = prom_con_obj.abacus_username
    password  = prom_con_obj.abacus_password
    return_dict={}
    for hostname in data[node_type]:
        return_dict[hostname] = {}
        return_dict[hostname]['storage'] = {}
        try:
            # ip_address = socket.gethostbyname(hostname)
            # return_dict[hostname]['lan_ip'] = ip_address
            # print(f"The IP address of {hostname} is {ip_address}")
            client = paramiko.SSHClient()
            client.load_system_host_keys() 
            client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            try:
                client.connect(hostname, port, username, password)
                commands = {"ram" : "free -g | awk '/Mem:/ {print $2}'" , "cores":"lscpu | awk '/^CPU\(s\):/ {print $2}'"}

                for label,command in commands.items():
                    stdin, stdout, stderr = client.exec_command(command)
                    out = stdout.read().decode('utf-8').strip()
                    if out and out!='':
                        return_dict[hostname][label] = out
                    else:
                        print(f"Unable to determine {label} value for {hostname}")
                
                storage_commands = {'root_partition':"df -h | awk '$6 == \"/\" {print $2}'",
                                    'kafka' : "df -h | awk '$6 == \"/data/kafka\" {print $2}'",
                                    'spark' : "df -h | awk '$6 == \"/data/spark\" {print $2}'",
                                    'dn1' : "df -h | awk '$6 == \"/data/dn1\" {print $2}'",
                                    'dn2' : "df -h | awk '$6 == \"/data/dn2\" {print $2}'",
                                    'dn3' : "df -h | awk '$6 == \"/data/dn3\" {print $2}'",
                                    'pg' : "df -h | awk '$6 == \"/pg\" {print $2}'",
                                    'data' : "df -h | awk '$6 == \"/data\" {print $2}'",
                                    'data_prometheus' : "df -h | awk '$6 == \"/data/prometheus\" {print $2}'",
                                    }

                for label,command in storage_commands.items():
                    stdin, stdout, stderr = client.exec_command(command)
                    out = stdout.read().decode('utf-8').strip()
                    if out and out!='':
                        return_dict[hostname]['storage'][label] = out
                    else:
                        print(f"Unable to determine {label} value for {hostname}")

            except Exception as e:
                print(f'ERROR : Unable connect to {hostname}' , e)

            finally:
                client.close()
        except socket.gaierror:
            print(f"Could not resolve {hostname}")

        if 'c2' in hostname:return_dict[hostname]['clst'] = "2"
        else:return_dict[hostname]['clst'] = "1"

    return return_dict

def extract_stack_details(nodes_file_path,prom_con_obj):
    with open(nodes_file_path,'r') as file:
        data = json.load(file)
    
    def extract_node_detail_wrapper(data, node_type, prom_con_obj):
        return extract_node_detail(data, node_type, prom_con_obj)

    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        future1 = executor.submit(extract_node_detail_wrapper, data, 'pnodes', prom_con_obj)
        future2 = executor.submit(extract_node_detail_wrapper, data, 'dnodes', prom_con_obj)
        future3 = executor.submit(extract_node_detail_wrapper, data, 'pgnodes', prom_con_obj)
        future4 = executor.submit(extract_node_detail_wrapper, data, 'monitoring_node', prom_con_obj)
        future5 = executor.submit(extract_node_detail_wrapper, data, 'other_nodes', prom_con_obj)

        completed_futures, _ = concurrent.futures.wait([future1, future2, future3, future4 , future5])

    pnodes = future1.result()
    dnodes = future2.result()
    pgnodes = future3.result()
    monitoring_node = future4.result()
    other_nodes = future5.result()

    data.update(pnodes)
    data.update(dnodes)
    data.update(pgnodes)
    data.update(monitoring_node)
    data.update(other_nodes)

    with open(nodes_file_path,'w') as file:
        json.dump(data,file,indent=4)

def add_test_env_details(nodes_file_path,doc):
    with open(nodes_file_path,'r') as file:
        data = json.load(file)
    table_dict={}
    table_dict={}
    table_dict['title'] = f"Test Environment details"
    table_dict['header'] =["node" ,'cluster', 'cores' , 'ram' , 'storage']
    table_dict['body']=[] 

    for node_type in ['pnodes' ,'dnodes', 'pgnodes' , 'monitoring_node']:
        for node in data[node_type]:
            curr_list=[]
            curr_list.append((node,0))
            curr_list.append((data[node]['clst'],0))
            curr_list.append((data[node]['cores'],0))
            curr_list.append((data[node]['ram'],0))
            storage_data=''
            for storage_type in data[node]['storage']:
                storage_data += storage_type + ':'
                storage_data += data[node]['storage'][storage_type]
                storage_data += ', '
            curr_list.append((storage_data,0))
            table_dict['body'].append(curr_list)
    
    doc = add_table(doc,table_dict)
    return doc

def push_data_to_mongo(load_name,json_path, prom_con_obj):
    try:
        connection_string=prom_con_obj.mongo_connection_string
        nodes_file_path = prom_con_obj.nodes_file_path
        client = pymongo.MongoClient(connection_string)
        db=client['all_loads']
        collection = db[load_name]
        with open(json_path,'r') as file:
            doc_to_insert=json.load(file)
        with open(nodes_file_path,'r') as file:
            test_env_details=json.load(file)
        doc_to_insert.update({"test_env_details":test_env_details})
        inserted_id = collection.insert_one(doc_to_insert).inserted_id
        print(f"Document pushed to mongo successfully into database {load_name} with id {inserted_id}")
    except:
        print(f"Document pushed to mongo successfully into database {load_name} with id {inserted_id}")